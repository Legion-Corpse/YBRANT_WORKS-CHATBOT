from __future__ import annotations

import ipaddress
import logging
import re
import socket
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol
from urllib.parse import urlsplit
from xml.etree.ElementTree import Element

import defusedxml.ElementTree as ET
import requests
import trafilatura
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

REQUEST_HEADERS = {"User-Agent": "Mozilla/5.0 (compatible; YbrantKnowledgeBot/1.0)"}
REQUEST_TIMEOUT = 30


class UnsafeUrlError(ValueError):
    """Raised when a URL targets a non-public / disallowed address (SSRF guard)."""


def assert_public_url(url: str) -> None:
    """Reject non-http(s) URLs and any host that resolves to a non-public IP.

    Defends the unauthenticated-adjacent ingest path against SSRF: an attacker
    could otherwise make the server fetch internal services or the cloud
    metadata endpoint (169.254.169.254). Every resolved address must be a
    global unicast address.
    """
    parts = urlsplit(url)
    if parts.scheme not in ("http", "https"):
        raise UnsafeUrlError(f"URL scheme not allowed: {parts.scheme or '(none)'}")
    host = parts.hostname
    if not host:
        raise UnsafeUrlError("URL has no host")
    try:
        infos = socket.getaddrinfo(host, parts.port or 0, proto=socket.IPPROTO_TCP)
    except socket.gaierror as exc:
        raise UnsafeUrlError(f"Cannot resolve host: {host}") from exc
    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if not ip.is_global or ip.is_reserved or ip.is_multicast:
            raise UnsafeUrlError(f"URL resolves to non-public address: {ip}")


def _local_xml_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1].lower()


def _xml_locs(root: Element) -> list[str]:
    locs: list[str] = []
    for node in root.iter():
        if _local_xml_name(node.tag) == "loc" and node.text:
            value = node.text.strip()
            if value:
                locs.append(value)
    return locs


@dataclass
class Document:
    source_id: str
    title: str
    url: str
    text: str


class Loader(Protocol):
    def can_load(self, target: str) -> bool: ...

    def load(self, target: str) -> list[Document]: ...


class WebLoader:
    def can_load(self, target: str) -> bool:
        return target.startswith(("http://", "https://"))

    MAX_REDIRECTS = 5

    def load(self, target: str) -> list[Document]:
        html = self._fetch_html(target)

        text = trafilatura.extract(html, include_comments=False) or ""
        soup = BeautifulSoup(html, "html.parser")
        if not text.strip():
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = re.sub(r"\n{3,}", "\n\n", soup.get_text(separator="\n")).strip()

        title = soup.title.get_text(strip=True) if soup.title else target
        if not text.strip():
            logger.warning("No content extracted from %s", target)
            return []
        return [Document(source_id=target, title=title, url=target, text=text)]

    @classmethod
    def _fetch_html(cls, url: str) -> str:
        """Return page HTML, JS-rendered when enabled, else static.

        Next.js renders some content client-side, so a static fetch misses it.
        When ``web_render_js`` is on, render via headless Chromium; on any render
        failure (Playwright/Chromium missing, navigation timeout) fall back to the
        SSRF-validated static fetch so ingest still works without the browser.
        """
        from app.config import settings

        if settings.web_render_js:
            try:
                return cls._render_html(url)
            except Exception:
                logger.warning(
                    "JS render failed for %s; falling back to static fetch", url
                )
        return cls._safe_get(url).text

    @staticmethod
    def _render_html(url: str) -> str:
        from playwright.sync_api import sync_playwright

        from app.config import settings

        # SSRF pre-check on the entry URL. (Playwright fetches sub-resources
        # itself, so per-hop validation isn't possible here; ingest is
        # admin-authenticated, which bounds the exposure — documented residual.)
        assert_public_url(url)
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            try:
                page = browser.new_page()
                page.goto(
                    url, wait_until="networkidle", timeout=settings.web_render_timeout_ms
                )
                return page.content()
            finally:
                browser.close()

    @classmethod
    def _safe_get(cls, url: str) -> requests.Response:
        """GET with SSRF validation on every hop (manual bounded redirects)."""
        current = url
        for _ in range(cls.MAX_REDIRECTS + 1):
            assert_public_url(current)
            resp = requests.get(
                current,
                headers=REQUEST_HEADERS,
                timeout=REQUEST_TIMEOUT,
                allow_redirects=False,
            )
            if resp.is_redirect or resp.is_permanent_redirect:
                location = resp.headers.get("location")
                if not location:
                    break
                current = requests.compat.urljoin(current, location)
                continue
            resp.raise_for_status()
            return resp
        raise UnsafeUrlError(f"Too many redirects starting from {url}")

    @staticmethod
    def discover_sitemap_urls(base_url: str) -> list[str]:
        sitemap_url = base_url.rstrip("/") + "/sitemap.xml"
        return WebLoader._discover_sitemap_urls(
            sitemap_url=sitemap_url,
            seen_sitemaps=set(),
            seen_pages=set(),
        )

    @staticmethod
    def _discover_sitemap_urls(
        sitemap_url: str,
        seen_sitemaps: set[str],
        seen_pages: set[str],
    ) -> list[str]:
        if sitemap_url in seen_sitemaps:
            return []
        seen_sitemaps.add(sitemap_url)

        # SSRF guard + bounded redirects on every sitemap fetch, same as page
        # fetches: a sitemapindex can name nested-sitemap URLs that we then fetch
        # recursively, so each must be validated as public. A fetch/parse failure
        # (network error, bad XML, private-IP target) skips that sitemap rather
        # than aborting the whole crawl.
        try:
            resp = WebLoader._safe_get(sitemap_url)
            root = ET.fromstring(resp.content)
        except Exception:
            logger.exception("Failed to read sitemap %s", sitemap_url)
            return []
        root_name = _local_xml_name(root.tag)
        locs = _xml_locs(root)

        if root_name == "sitemapindex":
            urls: list[str] = []
            for nested_sitemap in locs:
                urls.extend(
                    WebLoader._discover_sitemap_urls(
                        sitemap_url=nested_sitemap,
                        seen_sitemaps=seen_sitemaps,
                        seen_pages=seen_pages,
                    )
                )
            return urls

        urls = []
        for page_url in locs:
            if page_url not in seen_pages:
                seen_pages.add(page_url)
                urls.append(page_url)
        return urls


class PDFLoader:
    def can_load(self, target: str) -> bool:
        return target.lower().endswith(".pdf") and Path(target).exists()

    def load(self, target: str) -> list[Document]:
        from pypdf import PdfReader

        path = Path(target)
        reader = PdfReader(path)
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        title = (reader.metadata.title if reader.metadata else None) or path.stem
        return [
            Document(
                source_id=str(path.resolve()), title=title, url=path.name, text=text
            )
        ]


class DocxLoader:
    def can_load(self, target: str) -> bool:
        return target.lower().endswith(".docx") and Path(target).exists()

    def load(self, target: str) -> list[Document]:
        import docx

        path = Path(target)
        doc = docx.Document(str(path))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [
            Document(
                source_id=str(path.resolve()), title=path.stem, url=path.name, text=text
            )
        ]


class TextLoader:
    def can_load(self, target: str) -> bool:
        return target.lower().endswith((".txt", ".md")) and Path(target).exists()

    def load(self, target: str) -> list[Document]:
        path = Path(target)
        text = path.read_text(encoding="utf-8", errors="replace")
        return [
            Document(
                source_id=str(path.resolve()), title=path.stem, url=path.name, text=text
            )
        ]


ALL_LOADERS: list[Loader] = [WebLoader(), PDFLoader(), DocxLoader(), TextLoader()]


def find_loader(target: str) -> Loader:
    for loader in ALL_LOADERS:
        if loader.can_load(target):
            return loader
    raise ValueError(f"No loader available for: {target}")
